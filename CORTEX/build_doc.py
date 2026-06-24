# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ORANGE=RGBColor(0xC4,0x5A,0x1A); BLACK=RGBColor(0x1C,0x19,0x17)
GREY=RGBColor(0x6B,0x65,0x60); WHITE=RGBColor(0xFF,0xFF,0xFF)
CREAM="F5F2EE"; GREYBG="EDEAE4"; LINE="D8D2C8"

doc=Document()
# default font
st=doc.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(10.5); st.font.color.rgb=BLACK
sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
for m in ("top","bottom","left","right"): setattr(sec,f"{m}_margin",Inches(1))

def set_cell_bg(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),fill); tcPr.append(sh)
def set_cell_borders(cell):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right"):
        e=OxmlElement(f"w:{edge}"); e.set(qn("w:val"),"single"); e.set(qn("w:sz"),"4")
        e.set(qn("w:color"),LINE); b.append(e)
    tcPr.append(b)

def heading(text,level):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.name="Arial"; r.bold=True
    if level==1:
        r.font.size=Pt(16); r.font.color.rgb=ORANGE
        p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(8)
        pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"6")
        bot.set(qn("w:color"),"C45A1A"); pbdr.append(bot); pPr.append(pbdr)
        p.style=doc.styles["Heading 1"]
        for rr in p.runs: pass
    elif level==2:
        r.font.size=Pt(12.5); r.font.color.rgb=BLACK
        p.paragraph_format.space_before=Pt(11); p.paragraph_format.space_after=Pt(6)
        p.style=doc.styles["Heading 2"]
    else:
        r.font.size=Pt(11); r.font.color.rgb=ORANGE
        p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
        p.style=doc.styles["Heading 3"]
    # reset run after style assign
    p.clear() if False else None
    return p

# Because assigning p.style overrides runs, build headings via styles configured up-front instead:
def style_setup():
    for lvl,(size,color) in {1:(16,ORANGE),2:(12.5,BLACK),3:(11,ORANGE)}.items():
        s=doc.styles[f"Heading {lvl}"]; s.font.name="Arial"; s.font.bold=True
        s.font.size=Pt(size); s.font.color.rgb=color
        s.paragraph_format.space_before=Pt({1:16,2:11,3:8}[lvl])
        s.paragraph_format.space_after=Pt({1:8,2:6,3:4}[lvl])
style_setup()

def H(text,level=1):
    p=doc.add_paragraph(style=f"Heading {level}")
    r=p.add_run(text)
    if level==1:
        pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"6")
        bot.set(qn("w:color"),"C45A1A"); pbdr.append(bot); pPr.append(pbdr)
    return p

def para(parts, after=7, size=10.5, italic=False, color=None, align=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=1.18
    if align: p.alignment=align
    if isinstance(parts,str): parts=[(parts,False)]
    for item in parts:
        if isinstance(item,tuple): txt,bold=item
        else: txt,bold=item,False
        r=p.add_run(txt); r.font.name="Arial"; r.font.size=Pt(size); r.bold=bold
        r.italic=italic
        if color: r.font.color.rgb=color
    return p

def bullet(parts, numbered=False):
    style="List Number" if numbered else "List Bullet"
    p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(4)
    p.paragraph_format.line_spacing=1.15
    if isinstance(parts,str): parts=[(parts,False)]
    for item in parts:
        txt,bold=item if isinstance(item,tuple) else (item,False)
        r=p.add_run(txt); r.font.name="Arial"; r.font.size=Pt(10.5); r.bold=bold
    return p

def spacer(pt=6):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(pt); return p

def page_break():
    doc.add_page_break()

def table(headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.autofit=False; t.allow_autofit=False
    hdr=t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].width=Inches(widths[i]); set_cell_bg(hdr[i],"C45A1A"); set_cell_borders(hdr[i])
        cp=hdr[i].paragraphs[0]; cp.paragraph_format.space_after=Pt(0)
        r=cp.add_run(h); r.font.name="Arial"; r.font.bold=True; r.font.size=Pt(9); r.font.color.rgb=WHITE
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].width=Inches(widths[ci]); set_cell_borders(cells[ci])
            bg = GREYBG if ci==0 else (CREAM if ri%2==0 else None)
            if bg: set_cell_bg(cells[ci],bg)
            cp=cells[ci].paragraphs[0]; cp.paragraph_format.space_after=Pt(0)
            cp.paragraph_format.line_spacing=1.05
            r=cp.add_run(str(val)); r.font.name="Arial"; r.font.size=Pt(8.5)
            r.bold = (ci==0)
    return t

# ================= COVER =================
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(90); p.paragraph_format.space_after=Pt(0)
r=p.add_run("CORTEX"); r.font.name="Arial"; r.bold=True; r.font.size=Pt(54); r.font.color.rgb=ORANGE
p=doc.add_paragraph(); r=p.add_run("The Business Intelligence Operating System")
r.font.name="Arial"; r.font.size=Pt(17); r.font.color.rgb=BLACK
p.paragraph_format.space_after=Pt(4)
p=doc.add_paragraph(); r=p.add_run("Business Plan & Go-to-Market Strategy")
r.font.name="Arial"; r.font.size=Pt(13); r.font.color.rgb=GREY
pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"12"); bot.set(qn("w:space"),"8"); bot.set(qn("w:color"),"C45A1A")
pbdr.append(bot); pPr.append(pbdr); p.paragraph_format.space_after=Pt(20)
para([("Category: ",True),("BIOS — Business Intelligence Operating System (a new category)",False)])
para([("Builder: ",True),("Karim Azzaoui · Edge Graphics / winsym.ai · Kuala Lumpur, Malaysia",False)])
para([("Markets: ",True),("Malaysia / Southeast Asia (anchor) · Europe / DACH · Global Premium Enterprise",False)])
para([("Version: ",True),("1.0 · June 2026 · Confidential",False)])
spacer(14)
para([("“We don’t sell a chatbot. We install the brain a company runs on.”",False)], italic=True, color=ORANGE, size=13)
page_break()

# ================= 1. EXEC SUMMARY =================
H("1. Executive Summary")
para([("CORTEX is a governed digital brain that any organization installs to stop losing knowledge, guessing at decisions, and re-doing work. ",True),
("It is productized from the MIPP PIOS — a governed “leadership brain” built for a national political organization, a domain where a single fabricated fact causes real damage. The hardest, most defensible parts of that system — governance, grounded retrieval, memory, the knowledge graph, the maturity model and the verification layer — are already industry-neutral. Only the vocabulary is domain-specific. That is exactly what makes a scalable software company.",False)])
para([("The category is the wedge. ",True),("Generic AI tools answer questions. CORTEX remembers, reasons, decides, executes and improves — and never fabricates, because every claim carries a source, a confidence band and, for high-stakes output, a named human approver. We are not “ChatGPT for your company.” We name and own the category: BIOS — the Business Intelligence Operating System.",False)])
para([("The model is build-once-sell-many. ",True),("One universal platform (three pillars, eight layers, a five-level maturity ladder), configured per industry by reusable templates, delivered by a repeatable methodology (the CORTEX Method) that walks each client up the ladder. New verticals are new templates — not new engineering.",False)])
H("Headline numbers (base case)",3)
table(["Metric","Year 1","Year 2","Year 3"],
 [["Total revenue (USD)","$214K","$921K","$2.59M"],
  ["Recurring revenue (ARR)","$148K","$670K","$1.95M"],
  ["Gross margin","72%","73%","74%"],
  ["EBITDA","($25K)","$154K","$765K"],
  ["Active clients","16","76","243"],
  ["Exit-rate MRR","$12.4K","$55.9K","$162.8K"]],
 [2.7,1.9,1.9,1.9])
spacer(6)
para([("The ask. ",True),("Bootstrap the services-led pilot phase from winsym.ai cash flow plus a small (USD 150–300K) angel/grant round to fund the first engineers and the platform build, converting to a priced seed once 2–3 reference clients and a hardened vertical template are in place.",False)])

# ================= 2. MARKET ANALYSIS =================
H("2. Market Analysis")
para("CORTEX sits at the intersection of five fast-growing, currently-fragmented markets. The strategic insight: buyers are forced to assemble a brain from four or five tools that don’t share a knowledge graph, governance model or memory — and none is grounded enough to trust for consequential work. CORTEX collapses the stack.")
H("2.1 AI consulting providers",2)
para([("The Big Four (Accenture, Deloitte, McKinsey QuantumBlack, BCG X) and a long tail of boutiques sell AI strategy and bespoke builds. Expensive, slow, and they leave nothing reusable: every engagement re-learns the client and the deliverable is a deck, not a compounding asset. ",False),("Gap: ",True),("no productized, governed system the client keeps and grows. CORTEX delivers the consulting outcome as installed, owned software.",False)])
H("2.2 AI Operating Systems",2)
para([("Palantir Foundry + AIP is the closest analogue — an ontology-plus-AI operational layer — but priced and scoped for governments and the Fortune 500 (custom, six-to-seven-figure, land-and-expand) and heavy to implement. ",False),("Gap: ",True),("there is no Palantir for the SME and mid-market. CORTEX brings ontology + governance + grounded AI down to 1–2,500-person organizations at a price they can pay.",False)])
H("2.3 Knowledge management systems",2)
para([("Glean, Guru and the AI-search category make knowledge findable. The KM software market is ~USD 16–26 billion in 2026, growing 14–18% annually; AI-in-KM is growing ~25%+ CAGR. ",False),("Gap: ",True),("these tools retrieve and answer — they don’t run processes, score decisions, or act. They are Maturity Level 2. CORTEX uses Level-2 knowledge as the floor, not the ceiling.",False)])
H("2.4 AI agents",2)
para([("Salesforce Agentforce, Microsoft Copilot agents and a wave of frameworks promise autonomous action; ~80% of enterprises are deploying generative AI by 2026. But agents acting on ungoverned, unverified knowledge amplify mistakes at scale. ",False),("Gap: ",True),("nobody sells the governed substrate underneath the agents. CORTEX’s iron rule — never automate on top of unverified knowledge — is the discipline the agent hype is missing.",False)])
H("2.5 Internal GPT / “company ChatGPT” solutions",2)
para([("ChatGPT Enterprise, Claude for Teams/Enterprise and internal RAG builds give staff a private assistant — powerful but generic: no business knowledge graph, weak/optional governance, no permanent memory, no decision scoring, persistent hallucination risk. ",False),("Gap: ",True),("the trust layer. For clinics, law firms and enterprises — the buyers with the most to spend — a tool that can fabricate is unusable. CORTEX is built no-fabrication-first.",False)])
H("2.6 Opportunities & market gaps",2)
bullet([("The mid-market “Palantir gap.” ",True),("Governed, ontology-driven AI exists only at the top; 99% of organizations are priced out. The single biggest opening.",False)])
bullet([("The trust gap. ",True),("Regulated buyers cannot use tools that hallucinate. Grounded, sourced, human-approved output is a wedge generic AI can’t copy without rebuilding its core.",False)])
bullet([("The integration gap. ",True),("Buyers stitch together search + chat + CRM + automation. One governed brain spanning all eight functions is simpler and stickier.",False)])
bullet([("The Southeast Asia timing gap. ",True),("SEA’s AI market is expanding from ~USD 12B (2026) toward ~USD 80B (2031); Malaysia has ~68% professional AI adoption and a national “AI Nation 2030” push, yet a documented governance and data-readiness gap. High intent, low maturity, weak governance = the perfect entry market for a governed brain.",False)])
bullet([("The compounding-asset gap. ",True),("Every competitor sells a tool; CORTEX sells an asset that gets more valuable and more locked-in the longer it runs.",False)])

# ================= 3. COMPETITOR ANALYSIS =================
page_break()
H("3. Competitor Analysis")
para("No incumbent occupies CORTEX’s position: a governed, grounded, full-stack business brain priced for organizations below the enterprise mega-cap tier. Each competitor owns one slice; none owns the whole brain with trust built in. Pricing reflects publicly reported 2026 figures (mostly sales-quoted, not list).")
table(["Player","What it is","Indicative 2026 price","Gap vs. CORTEX"],
 [["Palantir Foundry/AIP","Ontology + AI operational layer","Custom; 6–7 figures, land-and-expand","Unreachable for SME/mid-market; heavy"],
  ["Glean","Enterprise AI search/assistant","~$50/user + $15 AI; ~100-seat min; ~$60K+ ACV","Answers only; no execution, decisions, memory graph"],
  ["Guru","AI knowledge management / wiki","~$25/seat; 10-seat minimum","Wiki, not a reasoning / operating brain"],
  ["Notion AI","Workspace + AI","Business ~$20/member; AI bundled","No governance, grounding or RBAC depth"],
  ["Microsoft 365 Copilot","AI across Office","$30/user enterprise add-on (on top of base)","Generic; no business knowledge graph or memory"],
  ["ChatGPT Enterprise","Private GPT for staff","~$50–60/user; 150-seat min; ~$108K/yr floor","No governed graph, memory or no-fabrication guarantee"],
  ["Claude for Teams / Enterprise","Private Claude for staff","~$30/user (Team); Enterprise custom","Assistant, not an operating system"],
  ["Salesforce Agentforce","Agents on CRM data","$2/conversation · $500/100K credits · ~$125/user","Locked to Salesforce; agents on ungoverned base"]],
 [1.5,1.95,2.35,2.6])
spacer(6)
H("3.1 Other relevant players",2)
para("Microsoft Copilot Studio and AWS/Google build-your-own stacks (you build the governance), vertical point solutions (Harvey for law, Abridge/Hippocratic for health), and regional SI / consulting shops. All either require you to build governance yourself or solve one vertical. CORTEX productizes the governance and spans verticals via templates.")
H("3.2 Why CORTEX wins (the moat)",2)
bullet([("Governance + trust layer. ",True),("Confidence bands, citations, human-approval tiers, no fabrication — what regulated buyers require, and a wedge generic AI can’t copy without rebuilding its core.",False)], numbered=True)
bullet([("Knowledge graph + memory. ",True),("Each client’s brain gets more valuable and more locked-in over time; switching means abandoning organizational memory.",False)], numbered=True)
bullet([("Template library. ",True),("A proprietary catalog of industry configurations makes every new vertical faster and cheaper than a competitor starting cold.",False)], numbered=True)
bullet([("The methodology. ",True),("The CORTEX Method turns implementation into a productized, teachable service — and the diagnostic is the sales pitch.",False)], numbered=True)

# ================= 4. OFFER STRUCTURE =================
page_break()
H("4. Offer Structure — The 4-Tier Ladder")
para("One platform, five maturity levels, four company sizes. Customers buy their way up the levels and switch on more layers — never re-platformed as they grow. The maturity gap itself is the value case in every sales conversation.")
table(["","STARTER","GROWTH","PROFESSIONAL","ENTERPRISE"],
 [["Target","Solo / micro (1–10)","SME (10–250)","Mid-market (250–2,500)","Enterprise (2.5K–10K+)"],
  ["Maturity","Level 1 → 2","Level 2 → 3","Level 3 → 4","Level 4 → 5"],
  ["Outcome","Find & trust answers","Brain runs the work","Brain helps decide","Business runs on it"],
  ["Layers active","2–3","4–5","6–7","All 8 lobes"],
  ["Documents","Up to 500","Up to 5,000","Up to 50,000","Unlimited"],
  ["Data sources","3 (uploaded)","8 (docs + first APIs)","15+ live integrations","Unlimited + custom"],
  ["AI functions","Grounded Q&A + sourcing","+ Content + ops drafting","+ Decision scoring + agents","+ Custom scoring, multi-brain"],
  ["Automations","Read-only briefs","Basic schedules","Workflows + integrations","Full cadence + triggers"],
  ["Dashboard","Search + ask box","+ Graph view, briefs","+ Scorecards, reports","+ Role-aware enterprise"],
  ["RBAC","Owner / flat","A few roles","Department clearances","Fine-grained, SSO, audit"],
  ["Deployment","Shared multi-tenant","Multi-tenant","MT or isolated","Single-tenant / VPC / on-prem"],
  ["Support","Self-serve + docs","Light onboarding","Guided implementation","Dedicated success + SLA"]],
 [1.35,1.65,1.65,1.75,1.8])
spacer(6)
para([("Three intertwined offers underneath the ladder: ",True),("(1) the Platform (recurring SaaS), (2) the Implementation (high-margin service that makes the software stick), and (3) the Industry Templates (reusable IP that makes each sale fast and bespoke-feeling). Add-ons monetize the variable layer: extra templates, additional brand-brains for agencies, premium integrations, compliance packs.",False)])

# ================= 5. PRICING =================
H("5. Pricing Strategy — Three Regional Price Books")
para("We price on value and maturity, not seats alone. Three commercial levers set price: maturity level reached, company size/usage, and deployment isolation. Southeast Asia is the anchor; Europe and Premium Enterprise are variants tuned to higher willingness-to-pay and stricter isolation needs.")
H("5.1 Malaysia / Southeast Asia (MYR) — anchor",2)
table(["","STARTER","GROWTH","PROFESSIONAL","ENTERPRISE"],
 [["Setup (one-time)","RM 6,000","RM 18,000","RM 45,000","RM 120,000+"],
  ["Monthly platform","RM 1,200","RM 3,500","RM 8,000","RM 18,000+"],
  ["Year-1 total","RM 20,400","RM 60,000","RM 141,000","RM 336,000+"]],
 [1.85,1.6,1.6,1.7,1.6])
spacer(5)
H("5.2 Europe / DACH (EUR)",2)
table(["","STARTER","GROWTH","PROFESSIONAL","ENTERPRISE"],
 [["Setup (one-time)","€ 2,500","€ 8,000","€ 22,000","€ 65,000+"],
  ["Monthly platform","€ 450","€ 1,400","€ 3,500","€ 8,500+"],
  ["Year-1 total","€ 7,900","€ 24,800","€ 64,000","€ 167,000+"]],
 [1.85,1.6,1.6,1.7,1.6])
spacer(5)
H("5.3 Global Premium Enterprise (USD) — single-tenant / VPC",2)
table(["","STARTER","GROWTH","PROFESSIONAL","ENTERPRISE"],
 [["Setup (one-time)","$ 3,000","$ 12,000","$ 35,000","$ 90,000+"],
  ["Monthly platform","$ 600","$ 2,000","$ 6,000","$ 15,000+"],
  ["Year-1 total","$ 10,200","$ 36,000","$ 107,000","$ 270,000+"]],
 [1.85,1.6,1.6,1.7,1.6])
spacer(5)
para([("Why this is defensible: ",True),("our Growth tier (~RM 60K / ~USD 13K Year-1) lands below the ~USD 60K–108K floors of Glean and ChatGPT Enterprise while delivering far more than search or chat. Our Enterprise tier still undercuts a Palantir engagement by an order of magnitude. We compete on value-per-dollar at the mid-market, not on per-seat price wars.",False)])

# ================= 6. MARGIN =================
H("6. Margin Analysis")
para("CORTEX has the classic services-to-product margin profile: implementation funds the early business at solid (40–48%) margins while subscription compounds at SaaS-grade (63–90%) margins. As templates mature and self-serve onboarding replaces hands-on capture, implementation hours fall and the subscription mix rises — pulling blended margin toward 80%+.")
H("6.1 Per-tier margins (SEA prices, USD)",2)
table(["","STARTER","GROWTH","PROFESSIONAL","ENTERPRISE"],
 [["Subscription gross margin","90%","88%","79%","63%"],
  ["Implementation gross margin","47%","45%","43%","47%"],
  ["Year-1 blended margin","78%","75%","68%","58%"]],
 [2.6,1.7,1.7,1.7,1.65])
spacer(5)
H("6.2 Cost structure",2)
bullet([("Subscription COGS: ",True),("LLM tokens + hosting (Postgres/pgvector, object store, optional OpenSearch/graph DB). ~$25/mo (Starter, shared) to ~$1,400/mo (Enterprise, single-tenant, heavy compute + governance).",False)])
bullet([("Implementation COGS: ",True),("consultant/engineer time running the CORTEX Method (16 hrs Starter → 320 hrs Enterprise at ~$42/hr internal cost in years 1–2). The cost line templates attack directly.",False)])
bullet([("OPEX: ",True),("team, S&M, G&A — $180K (Y1) → $520K (Y2) → $1.15M (Y3), crossing into positive EBITDA in Year 2.",False)])
H("6.3 Scalability",2)
para("The core architecture is invariant from 1 to 10,000 employees — only how much is switched on, how isolated the deployment is, and how granular the governance is flex with size. A stateless API layer scales horizontally; Postgres+pgvector carries the long tail; the LLM layer is provider-pluggable and cost-tiered. Commercially, the same platform sells across five maturity stages, four sizes and many templates — growth is configuration, not migration, and governance becomes more premium (not more costly to us) as clients grow.")
H("6.4 Personnel requirement",2)
table(["Function","Year 1","Year 2","Year 3"],
 [["AI / full-stack engineering","1","3","6"],
  ["Implementation consulting","1","3","7"],
  ["Sales / partnerships","0","2","5"],
  ["Customer success","0","1","4"],
  ["Marketing / content","1","1","3"],
  ["Founder + ops / admin","1","2","3"],
  ["Total headcount","4","12","28"]],
 [3.6,1.9,1.9,1.9])

# ================= 7. WEBSITE =================
page_break()
H("7. Website Structure & Copy")
para([("Six sections built to convert serious buyers: lead with the category and the trust promise, prove it with the maturity diagnostic, route everyone to a booked demo. Structure: ",False),("Home · Solutions · Industries · Pricing · Case Studies · Book Demo.",True)])
H("Home",2)
para([("Hero headline: ",True),("“Your company already has a brain. It’s just scattered across people, files and apps.”",False)])
para([("Sub-headline: ",True),("CORTEX is the governed AI operating system that captures what your business knows, runs how it works, and helps it decide — without ever making things up.",False)])
para([("CTAs: ",True),("Primary — Book a demo.  Secondary — Take the 2-minute Maturity Assessment.",False)])
para([("Value props: ",True),("(1) Never lose knowledge again — one trusted source of truth. (2) Decisions on evidence, not opinion — every answer sourced and confidence-rated. (3) Built for trust — no fabrication, full access control, human approval for high-stakes output.",False)])
H("Solutions",2)
para("Organized by the eight brain layers framed as outcomes: Knowledge (“end ask-the-veteran”), Operations (“meetings become assigned work”), Marketing & Content (“on-brand at scale”), Sales (“every rep performs like your best rep”), Customer (“one memory of every relationship”), Decision (“score high-stakes choices”), Intelligence & Risk (“live radar + crisis playbook”), Automation (“the business runs on a cadence”). Each labelled with the maturity level it unlocks.")
para([("Value proposition: ",True),("“One brain, eight functions, zero hallucinations.”",False)])
H("Industries",2)
para("A page per hardened template — Law Firms, Medical Practices, Hotels & Hospitality, Real Estate, Marketing Agencies, Restaurants — each showing the layers that lead, the compliance red-lines we respect (attorney-client privilege, HIPAA/PDPA, food-safety), and a vertical-specific demo. The trust layer is the hero on regulated pages.")
para([("Value proposition: ",True),("“We don’t adapt your business to software. We configure the brain to your industry.”",False)])
H("Pricing",2)
para("The 4-tier ladder shown as a maturity journey (Starter → Enterprise), region-aware (MYR/EUR/USD), with the maturity diagnostic embedded: “Where is your business today?” Anchored on value vs. the cost of lost knowledge and bad decisions, not per-seat comparison. CTA on every tier.")
H("Case Studies",2)
para("Proof over promises. Each study: the maturity gap before → what we captured and switched on → measurable result (onboarding time cut, questions deflected, decision quality, hours saved). Lead with the PIOS lineage as credibility anchor (“built where one fabricated fact causes real damage”), then add commercial references as they land.")
H("Book Demo",2)
para("A single frictionless booking flow. The demo is the differentiator: a live, grounded, fully-sourced answer from a sample brain — then a deliberate “I don’t know” on an out-of-scope question. We sell the governance, not the magic.")

# ================= 8. SALES =================
H("8. Sales Strategy")
H("8.1 Target customers & industries",2)
bullet([("Beachhead verticals (regulated = trust-layer sells itself): ",True),("law firms, medical/dental practices, then hotels, real estate, marketing agencies, multi-location hospitality.",False)])
bullet([("Buyer: ",True),("owner / managing partner / GM of a 10–250-person SME, and department heads in 250–2,500-person mid-market firms — non-technical, ROI-driven, feeling the pain of lost knowledge and inconsistent service.",False)])
bullet([("Geography: ",True),("Kuala Lumpur / Malaysia first (warm network, winsym.ai presence, Google + Meta Partner credibility), then SEA, then DACH via the German-language advantage.",False)])
H("8.2 Acquisition process",2)
para([("The diagnostic is the pitch. ",True),("Every conversation opens by placing the prospect on the maturity ladder: “You’re operating at Level 1 but trying to make Level 4 decisions.” The gap is the value case. Funnel: content/SEO + warm outreach + speaking → free Maturity Assessment → discovery call → grounded live demo → paid pilot (one or two layers) → land at Level 2.",False)])
H("8.3 Demo process",2)
para("Pre-load a small sample brain for the prospect’s vertical. Show a grounded, sourced answer; then deliberately ask something not in the knowledge base to prove it says “I don’t know” and routes the gap rather than fabricating. That refusal to make things up is what closes regulated buyers.")
H("8.4 Upselling / land-and-expand",2)
para("Land small on the one or two layers that hurt most (usually Knowledge + Operations), prove ROI fast at Level 2, then sell the climb: more layers, higher maturity, more departments, deeper integrations, additional templates/brand-brains. Net revenue retention modeled at 115% — the brain’s growing lock-in does the selling. Win 2–3 reference clients in one vertical, harden the template, then repeat.")

# ================= 9. ROADMAP =================
page_break()
H("9. Product Roadmap")
table(["Phase","Theme","What we build / do","Outcome"],
 [["Phase 1","Done-for-you Brain Setup","Services-led pilots for 2–3 clients in one vertical using the file-based/PoC approach + the CORTEX Method. Learn the template.","Revenue + references; first hardened template"],
  ["Phase 2","Brain as a Service","Productize the platform: graph + RBAC + grounded RAG + trust layer first, then dashboard and user layers.","Recurring SaaS revenue; repeatable onboarding"],
  ["Phase 3","Managed AI Operating System","Switch on the Automation/Cadence layer and live integrations for clients who’ve earned the maturity. Read-only automations first.","Level-5 clients; higher ACV; deep lock-in"],
  ["Phase 4","Own SaaS Platform","Lower the floor (solo/SME self-serve) and raise the ceiling (enterprise single-tenant), with an expanding template marketplace.","Scalable self-serve + enterprise; template network effects"]],
 [1.3,2.0,3.5,2.0])
spacer(6)
para([("Sequencing logic: ",True),("services-led first to learn the verticals and fund the build (a brain must be trusted before it’s bought), then convert to platform-led recurring revenue as templates mature. Never automate (Phase 3) before the knowledge it acts on is verified — the iron rule of the maturity model is also the rule of the roadmap.",False)])

# ================= 10. FUNDING =================
H("10. Funding Strategy")
H("10.1 Bootstrapping (now)",2)
para("Fund Phase 1 from winsym.ai consulting cash flow. Services revenue from the first done-for-you brains directly finances the platform build — the capital-efficient services-to-product path. Keep the team lean (4 in Year 1) and let implementation margin cover burn.")
H("10.2 Grants & support programmes",2)
para("Malaysia’s “AI Nation 2030” agenda and MSME AI-readiness initiatives (MDEC, Cradle, MyDIGITAL, plus Microsoft/partner programmes), and German/EU digitalization grants for the DACH expansion. Non-dilutive capital plus government credibility — valuable for selling to public-sector and regulated buyers.")
H("10.3 Angel investors",2)
para("A small USD 150–300K angel round once 2–3 reference clients and one hardened template exist — targeting SEA operator-angels and ex-founders who can open enterprise doors. Use of funds: first 2 engineers + a sales hire + the platform MVP (Phase 2).")
H("10.4 Venture capital",2)
para("A priced seed (~USD 1.5–3M) once the platform shows recurring revenue, net revenue retention >110%, and a repeatable template motion. The pitch is category creation (BIOS) + a defensible governance moat + a services-to-SaaS margin story + the SEA-to-global expansion path. Regional funds for the seed; global funds at Series A.")
H("10.5 Strategic partners",2)
para("Cloud and model providers (Anthropic, Microsoft, AWS, Google — co-sell and credits), systems integrators and vertical associations (bar councils, medical/hospitality bodies) for distribution, and Google/Meta Partner status as trust signals. Strategic capital later can come from a cloud/enterprise-software player seeking a governed-AI layer for the mid-market.")

# ================= 11. CONCLUSION =================
page_break()
H("11. Conclusion — The Company You Can Sell")
para([("CORTEX builds the Business Intelligence Operating System: a governed digital brain that any organization installs to stop losing knowledge, guessing at decisions, and re-doing work. ",True),("We sell one platform — configured per industry by reusable templates, delivered via a methodology that walks each client up a five-level maturity ladder — from a productivity tool for a solo founder to the operating system of a 10,000-person enterprise.",False)])
para([("The positioning, in one line: ",True),("the rigor of McKinsey, the data-platform discipline of Palantir, and the everyday usability of Notion — fused into one trustworthy product, priced for the 99% of organizations the incumbents ignore.",False)])
H("Revenue potential",2)
para("The base case reaches ~USD 2.6M revenue and ~USD 1.95M ARR by Year 3 at ~74% gross margin and ~29% EBITDA margin, with 243 active clients and an exit-rate MRR of ~$163K — a credible, capital-efficient path to a $10M+ ARR business. Because the platform never re-platforms a customer as they grow, and governance becomes more premium with scale, the model’s defensibility and expansion revenue increase over time. The ceiling is set by how many verticals we templatize and how far up the maturity ladder we take each client.")
para([("The next move: ",True),("close the first 2–3 done-for-you brains in a single beachhead vertical, harden that template, and use the references to raise the angel round that funds the platform. The framework is done. The company is a sequencing problem now — and the sequence is clear.",False)])
spacer(8)
para([("CORTEX — Winning systems, backed by a brain.",False)], italic=True, color=ORANGE, size=11.5)
para([("Figures are illustrative model anchors built on the attached financial model; validate against live prospects before committing. Competitor pricing reflects publicly reported 2026 figures and is mostly sales-quoted, not list.",False)], color=GREY, size=8)

# Footer
ftr=sec.footer; fp=ftr.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=fp.add_run("CORTEX — Confidential Business Plan"); fr.font.size=Pt(8); fr.font.color.rgb=GREY; fr.font.name="Arial"

doc.save("/sessions/brave-zen-fermat/mnt/Winsym Coaching/CORTEX/CORTEX_Business_Plan.docx")
print("doc saved")
